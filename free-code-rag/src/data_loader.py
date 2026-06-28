"""
Load documents from the knowledge base directory.
Supports PDF, DOCX, TXT, and Markdown (.md) file types.
Default directory: ~/.free-code/knowledgeBase/
"""
import os
from pathlib import Path
from typing import List, Optional

from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document

KNOWLEDGE_BASE_DIR = os.path.join(os.path.expanduser("~"), ".free-code", "knowledgeBase")

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".doc"}

# Sidecar metadata for documents; kept on disk but excluded from vector indexing and /query.
KNOWLEDGE_METADATA_SUFFIX = ".knowledge.md"


def is_knowledge_metadata_filename(name: str) -> bool:
    """True if basename is a paired metadata file (e.g. doc1.knowledge.md)."""
    return name.casefold().endswith(KNOWLEDGE_METADATA_SUFFIX)


def is_knowledge_metadata_source(source: str) -> bool:
    """True if a chunk metadata 'source' path refers to a *.knowledge.md file."""
    if not source or not str(source).strip():
        return False
    s = str(source).strip()
    if is_knowledge_metadata_filename(Path(s).name):
        return True
    # Legacy or odd loaders: full path may not end with .name as expected
    return s.casefold().endswith(KNOWLEDGE_METADATA_SUFFIX) or (
        KNOWLEDGE_METADATA_SUFFIX in s.casefold() and Path(s).name.casefold().endswith(KNOWLEDGE_METADATA_SUFFIX)
    )


def is_knowledge_sidecar_chunk_content(text: str) -> bool:
    """Heuristic for legacy indexes that omitted 'source' in metadata but embedded *.knowledge.md text."""
    head = (text or "").strip()[:800]
    if not head.startswith("#"):
        return False
    return "knowledge base metadata" in head.lower()


def get_knowledge_base_dir() -> Path:
    path = Path(KNOWLEDGE_BASE_DIR)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _sanitize_data_dir(data_dir: str) -> Path:
    if ".." in data_dir:
        raise ValueError("Path traversal (..) is not allowed in data_dir")
    path = Path(data_dir).resolve()
    if not path.exists():
        return path
    if not path.is_dir():
        raise ValueError(f"data_dir must be a directory: {path}")
    return path


def _load_pdf(file_path: Path) -> List[Document]:
    try:
        import fitz
        doc = fitz.open(str(file_path))
        documents = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text().strip()
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": str(file_path), "page": page_num},
                    )
                )
        doc.close()
        return documents
    except ImportError:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        documents = []
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                documents.append(
                    Document(
                        page_content=text.strip(),
                        metadata={"source": str(file_path), "page": page_num},
                    )
                )
        return documents


def _load_docx(file_path: Path) -> List[Document]:
    from docx import Document as DocxDocument
    docx_doc = DocxDocument(str(file_path))
    full_text = "\n".join(p.text for p in docx_doc.paragraphs if p.text.strip())
    if not full_text.strip():
        return []
    return [
        Document(
            page_content=full_text,
            metadata={"source": str(file_path)},
        )
    ]


def _load_txt(file_path: Path) -> List[Document]:
    try:
        loader = TextLoader(str(file_path))
        return loader.load()
    except Exception as e:
        print(f"[ERROR] Failed to load TXT {file_path}: {e}")
        return []


def load_file(file_path: Path) -> List[Document]:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _load_docx(file_path)
    elif ext in (".txt", ".md"):
        return _load_txt(file_path)
    else:
        print(f"[WARN] Unsupported file type: {ext} ({file_path})")
        return []


def load_all_documents(data_dir: Optional[str] = None) -> List[Document]:
    data_path = _sanitize_data_dir(data_dir) if data_dir else get_knowledge_base_dir()
    documents: List[Document] = []

    if not data_path.exists():
        return documents

    for ext in SUPPORTED_EXTENSIONS:
        for file_path in data_path.glob(f"**/*{ext}"):
            if is_knowledge_metadata_filename(file_path.name):
                continue
            docs = load_file(file_path)
            documents.extend(docs)

    return documents


def load_new_documents(
    data_dir: str, indexed_sources: List[str]
) -> List[Document]:
    exclude = {str(Path(s).resolve()) for s in indexed_sources if s}
    data_path = _sanitize_data_dir(data_dir)
    documents: List[Document] = []

    if not data_path.exists():
        return documents

    for ext in SUPPORTED_EXTENSIONS:
        for file_path in data_path.glob(f"**/*{ext}"):
            if is_knowledge_metadata_filename(file_path.name):
                continue
            resolved = str(file_path.resolve())
            if resolved in exclude:
                continue
            docs = load_file(file_path)
            documents.extend(docs)

    return documents


def list_knowledge_base_files(data_dir: Optional[str] = None) -> List[dict]:
    data_path = Path(data_dir) if data_dir else get_knowledge_base_dir()
    if not data_path.exists():
        return []

    files = []
    for ext in SUPPORTED_EXTENSIONS:
        for file_path in data_path.glob(f"**/*{ext}"):
            stat = file_path.stat()
            files.append({
                "name": file_path.name,
                "path": str(file_path),
                "size_bytes": stat.st_size,
                "modified": stat.st_mtime,
            })

    return sorted(files, key=lambda f: f["name"])
